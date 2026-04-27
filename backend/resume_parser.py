"""Parse PDF and DOCX resume uploads into plain text."""
import io
from typing import Tuple

from PyPDF2 import PdfReader
import docx


def parse_resume_bytes(filename: str, data: bytes) -> Tuple[str, str]:
    """Return (plain_text, detected_format). Raises ValueError on unsupported."""
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return _parse_pdf(data), "pdf"
    if lower.endswith(".docx"):
        return _parse_docx(data), "docx"
    if lower.endswith(".doc"):
        # old binary .doc not supported -> instruct user to save as docx
        raise ValueError("Legacy .doc format not supported. Please upload a .docx or .pdf file.")
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="ignore"), "txt"
    raise ValueError("Only PDF, DOCX, or TXT resumes are supported.")


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(chunks).strip()
    if not text:
        raise ValueError("Could not extract text from the PDF. The file might be scanned/empty.")
    return text


def _parse_docx(data: bytes) -> str:
    d = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in d.paragraphs if p.text]
    # include tables cells too
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("Could not extract text from the DOCX file.")
    return text
